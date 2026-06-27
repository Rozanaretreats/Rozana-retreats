"""Generate Rozana Flow Inventory v2 Word document."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCS = ROOT / "docs" / "Rozana-Flow-Inventory-v2.docx"
OUT_DOWNLOADS = Path(r"c:\Users\viswa\Downloads\Rozana resorts\Rozana-Flow-Inventory-v2-roles.docx")


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement

    shading = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shading.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D8E8DF") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("FLOW INVENTORY")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Rozana Resorts — Attendance, Leave & Housekeeping (v2)")
    sr.font.size = Pt(12)
    sr.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        "Supersedes: Rozana-Flow-Inventory.docx (23 Jun 2026)\n"
        "Issued: 20 Jun 2026  ·  Personas: roles only (no named individuals)"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Legend: ").bold = True
    p.add_run("✅ Built  ·  🔶 Partial  ·  ⏳ Go-live pending  ·  🔮 Discussed, not built")

    add_heading(doc, "Summary", 2)
    add_table(
        doc,
        ["Metric", "v1 docx (Jun 23)", "v2 build (now)"],
        [
            ["Flows defined", "16 (F1–F16)", "24 (+8 new)"],
            ["Staff app", "Deferred", "✅ Live — My tasks / attendance / leave"],
            ["Photo proof", "Before + after pair", "✅ Checklist + 2 random verification photos"],
            ["WhatsApp EOD (F14)", "In scope", "⏳ UI ready; Edge Function not wired"],
            ["Face / biometric enrol (F15)", "In scope", "🔶 Roster + login only; no face templates"],
            ["Built in app", "Planning stage", "21 ✅ · 2 🔶 · 1 ⏳"],
        ],
    )

    add_heading(doc, "Module map", 2)
    add_table(
        doc,
        ["Module", "Flow prefix", "Flows"],
        [
            ["Module 1 — Staff Attendance", "AT", "AT01–AT07"],
            ["Module 2 — Leave", "LV", "LV01–LV05"],
            ["Module 3 — Housekeeping", "HK", "HK01–HK08"],
            ["Module 4 — Owner report & alerts", "RP", "RP01–RP04"],
            ["Platform", "PL", "PL01–PL08, PL15"],
        ],
    )

    add_heading(doc, "F1–F16 mapping (original docx → current build)", 2)
    add_table(
        doc,
        ["Old ID", "Job to be done", "New ID(s)", "Status", "Implementation"],
        [
            ["F1", "Check-in on arrival", "AT01", "✅", "Biometric → attendance_punches; Attendance page"],
            ["F2", "Check-out; hours close for the day", "AT02, AT04", "✅", "Punch log + EOD rollup on Reports"],
            ["F3", "Punches → worked hours per person", "AT04", "✅", "lib/attendance.ts; Reports EOD block"],
            ["F4", "Catch anomalies same day", "AT03, AT06, RP01", "🔶", "Device anti-fraud; operations manager mark absent; no auto double-punch flag yet"],
            ["F5", "Record staff leave", "LV01, LV04", "✅", "Operations manager Leave page; housekeeping staff My leave request"],
            ["F6", "Leave day ≠ false absence", "LV02", "✅", "staffStatus.ts on-leave vs absent"],
            ["F7", "Availability before room assign", "LV03", "✅", "Housekeeping Assign — present staff only"],
            ["F8", "Assign room to named staff", "HK01", "✅", "Housekeeping Assign tab"],
            ["F9", "Staff does room + photos, low friction", "HK04, HK05, PL04", "✅", "Housekeeping staff portal: checklist → 2 verification photos"],
            ["F10", "Done only when genuinely proven", "HK05", "✅", "Cannot complete without verification photos"],
            ["F11", "Skipped / unproven rooms stay visible", "HK03, RP01", "✅", "Tasks board + red alert on Reports"],
            ["F12", "EOD aggregate per property", "RP02", "✅", "Reports per-property preview"],
            ["F13", "Same-day incomplete alert", "RP01", "✅", "Outstanding rooms on owner report"],
            ["F14", "WhatsApp to resort owners, untouched by operations manager", "RP03", "⏳", "Send button disabled until Edge Function"],
            ["F15", "Enrol staff identity / photo once", "PL03, PL06, PL15", "🔶", "Team + staff_logins yes; face template no"],
            ["F16", "Role scope; operations manager cannot weaken controls", "PL02, RP04", "✅", "Route guards; operations manager no Reports; punches read-only"],
        ],
    )

    add_heading(doc, "New flows (beyond original inventory)", 2)
    add_table(
        doc,
        ["ID", "Flow", "Status", "Where in Rozana Ops"],
        [
            ["LV04", "Staff submits leave request (pending)", "✅", "StaffMyLeavePage → leave_records"],
            ["LV05", "Operations manager approves / rejects leave", "✅", "LeavePage; OpsContext.approveLeave"],
            ["HK04", "In-room cleaning checklist (6 items)", "✅", "CleaningChecklistModal; cleaning_checklist JSONB"],
            ["HK05", "2 random verification photos (bathroom excluded)", "✅", "VerificationPhotoModal; hk-photos bucket"],
            ["HK06", "Cleaning start / finish timestamps", "✅", "cleaning_started_at, cleaning_finished_at"],
            ["HK08", "Operations manager task monitor + photo gallery", "✅", "HousekeepingPage Tasks + HK photo proof"],
            ["PL06", "HK credential provisioning (Team tab)", "✅", "Housekeeping Team; staff_logins"],
            ["PL08", "Private photo storage + signed URLs", "✅", "lib/hkPhotos.ts; hk-photos bucket"],
            ["AT06", "Operations manager marks absent after shift start", "✅", "AttendancePage; attendance_absences"],
            ["AT07", "Manual punch testing (pre-biometric)", "🔶", "VITE_ALLOW_MANUAL_PUNCHES only"],
        ],
    )

    add_heading(doc, "Module 1 — Attendance", 2)
    add_table(
        doc,
        ["ID", "Flow", "Personas", "Status", "Success criteria"],
        [
            ["AT01", "Staff check-in", "Housekeeping staff · System", "✅", "Identity, on-site presence, timestamp captured"],
            ["AT02", "Staff check-out", "Housekeeping staff · System", "✅", "Check-out paired to check-in; hours calculable"],
            ["AT03", "Buddy-punch / off-site rejection", "System · Device", "✅", "Enforced on biometric hardware"],
            ["AT04", "Daily attendance rollup", "System · Owner", "✅", "EOD attendance block on Reports"],
            ["AT05", "Capture method / shift setup", "Owner · Operations manager", "✅", "Shift times on properties table"],
            ["AT06", "Mark absent (after shift)", "Operations manager · Owner", "✅", "Absent with reason; not editable punch"],
            ["AT07", "Manual test punches", "Operations manager (dev)", "🔶", "Pre-go-live testing only"],
        ],
    )

    add_heading(doc, "Module 2 — Leave", 2)
    add_table(
        doc,
        ["ID", "Flow", "Personas", "Status", "Success criteria"],
        [
            ["LV01", "Leave recorded", "Operations manager · Owner", "✅", "Leave saved for staff + date range"],
            ["LV02", "On-leave in availability", "System", "✅", "On-leave, not absent, on leave days"],
            ["LV03", "Availability before HK assign", "Operations manager", "✅", "Only present staff assignable"],
            ["LV04", "Housekeeping staff leave request", "Housekeeping staff", "✅", "Pending request visible to operations manager"],
            ["LV05", "Approve / reject leave", "Operations manager · Owner", "✅", "Status + reviewed_by recorded"],
        ],
    )

    add_heading(doc, "Module 3 — Housekeeping", 2)
    add_table(
        doc,
        ["ID", "Flow", "Personas", "Status", "Success criteria"],
        [
            ["HK01", "Assign room to staff", "Operations manager", "✅", "One staff per room; visible in housekeeping staff app"],
            ["HK02", "todo → cleaning → done", "Housekeeping staff · Operations manager", "✅", "Status tracked end-to-end"],
            ["HK03", "Open tasks on owner reports", "System", "✅", "Outstanding rooms in EOD"],
            ["HK04", "Cleaning checklist", "Housekeeping staff", "✅", "6 items; all required before photos"],
            ["HK05", "Verification photos", "Housekeeping staff", "✅", "2 random items; room cannot complete without"],
            ["HK06", "Cleaning timestamps", "System", "✅", "Start and finish times stored"],
            ["HK08", "Operations manager monitor + gallery", "Operations manager · Owner", "✅", "Live progress + proof review"],
        ],
    )

    add_heading(doc, "HK cleaning sub-flow (housekeeping staff app)", 2)
    doc.add_paragraph(
        "1. Operations manager assigns room (HK01) → status: To do\n"
        "2. Housekeeping staff taps Start cleaning (HK04) → status: In progress; checklist opens\n"
        "3. Housekeeping staff completes all checklist items (mop, bedsheet, blanket, amenities, bathroom, trash)\n"
        "4. System picks 2 random completed items for verification (bathroom never selected)\n"
        "5. Housekeeping staff captures 2 verification photos (HK05) → status: Done (HK06 timestamps saved)\n"
        "6. Operations manager reviews on Housekeeping Tasks + photo gallery (HK08)"
    )

    add_heading(doc, "Module 4 — Owner report & alerts", 2)
    add_table(
        doc,
        ["ID", "Flow", "Personas", "Status", "Success criteria"],
        [
            ["RP01", "Incomplete task alert (same day)", "System · Owner", "✅", "Unproven rooms flagged red"],
            ["RP02", "EOD compile attendance + HK", "System", "✅", "Per-property summary with traceability"],
            ["RP03", "WhatsApp to resort owners", "System", "⏳", "Resort owners receive proof-backed summary (not via operations manager)"],
            ["RP04", "Operations manager cannot edit report", "System", "✅", "Operations manager blocked from Reports route"],
        ],
    )

    add_heading(doc, "Platform", 2)
    add_table(
        doc,
        ["ID", "Flow", "Personas", "Status", "Success criteria"],
        [
            ["PL01", "Supabase schema + RLS", "Build", "✅", "Migrations 001–016; demo RLS policies"],
            ["PL02", "Auth & roles", "All roles", "🔶", "Housekeeping staff in Supabase; owner / operations manager demo logins"],
            ["PL03", "Roster + room inventory", "Owner · Operations manager", "✅", "32 rooms; Team tab roster"],
            ["PL04", "Housekeeping staff portal (mobile v1)", "Housekeeping staff", "✅", "/my-tasks, /my-attendance, /my-leave"],
            ["PL06", "Staff login provisioning", "Operations manager", "✅", "Team tab email + password → staff_logins"],
            ["PL08", "Private HK photos", "System", "✅", "hk-photos bucket; signed URLs"],
            ["PL15", "Face / biometric enrolment", "Operations manager · Device", "🔮", "Kiosk + face templates — not built"],
        ],
    )

    add_heading(doc, "Screens by role", 2)
    add_table(
        doc,
        ["Screen / route", "Owner", "Operations manager", "Housekeeping staff"],
        [
            ["Login", "✓", "✓", "✓ (credentials from Team tab)"],
            ["Attendance", "✓ all properties", "✓ Ooty only", "—"],
            ["Leave", "✓", "✓", "—"],
            ["Housekeeping", "✓", "✓", "—"],
            ["Reports / EOD", "✓", "✗", "—"],
            ["Property toggle (Both)", "✓", "✗", "—"],
            ["My tasks", "—", "—", "✓"],
            ["My attendance (read-only)", "—", "—", "✓"],
            ["My leave", "—", "—", "✓"],
        ],
    )

    add_heading(doc, "Permission matrix", 2)
    add_table(
        doc,
        ["Action", "Operations manager", "Owner", "Housekeeping staff"],
        [
            ["View attendance", "Ooty", "All properties", "Own punches only"],
            ["Record / approve leave", "✓", "✓", "Request only"],
            ["Mark absent", "✓", "✓", "✗"],
            ["Assign HK + monitor", "✓ Ooty", "✓", "✗"],
            ["Complete cleaning + photos", "✗", "✗", "✓ assigned rooms"],
            ["Edit biometric punch", "✗", "✗", "✗"],
            ["Reports / EOD / WhatsApp", "✗", "✓", "✗"],
            ["Property toggle", "✗", "✓", "✗"],
        ],
    )

    add_heading(doc, "Daily operations loop (v2)", 2)
    doc.add_paragraph(
        "Morning: Housekeeping staff punch in on biometric device (AT01) → Attendance shows Present / Not in yet.\n"
        "Operations manager checks availability (LV03) → assigns rooms on Housekeeping (HK01).\n"
        "Housekeeping staff open My tasks on phone → Start cleaning → checklist → 2 verification photos → Done (HK04–HK06).\n"
        "Operations manager monitors Tasks tab and photo gallery (HK08).\n"
        "Evening: Owner opens Reports → EOD preview per property (RP02); outstanding rooms flagged (RP01).\n"
        "Go-live: WhatsApp to resort owners (RP03) — pending Edge Function."
    )

    add_heading(doc, "Build order (as implemented)", 2)
    doc.add_paragraph(
        "PL01 Schema → PL02 Roles → PL03 Rooms/roster → PL06 Staff logins → PL04 Staff portal → "
        "HK04 Checklist + HK05 Photos + PL08 Storage → AT/LV manager flows → HK01 Assign → "
        "RP01–RP02 Reports → RP03 WhatsApp (pending) → PL15 Face kiosk (future)."
    )

    add_heading(doc, "Explicitly out of scope / future", 2)
    add_table(
        doc,
        ["Area", "Notes"],
        [
            ["Booking / check-in / walk-in", "Deferred — eZee replacement phase"],
            ["Cash log, OTA, maintenance", "Not in build brief"],
            ["Payroll", "Explicit exclusion"],
            ["CCTV hardware", "Client responsibility"],
            ["Production auth for owner / operations manager", "Move to Supabase Auth + profiles table"],
            ["PWA / App Store install", "Mobile web works; installable app optional"],
            ["Room rename UI", "SQL update on rooms table today"],
            ["Face recognition kiosk", "Discussed; replaces/supplements biometric at property"],
        ],
    )

    add_heading(doc, "Role definitions", 2)
    doc.add_paragraph(
        "Owner — Full property scope (toggle Ooty · Kannur · Both); Attendance, Leave, Housekeeping, "
        "Reports / EOD. Cannot edit biometric punches.\n\n"
        "Operations manager — On-site role for assigned property (e.g. Ooty only); Attendance, Leave, "
        "Housekeeping. Cannot access Reports or edit punches. Subject to the same attendance and report "
        "controls as staff (cannot weaken owner-facing proof).\n\n"
        "Housekeeping staff — My tasks, My attendance (read-only), My leave (request). Completes assigned "
        "rooms via checklist and verification photos. Cannot see other staff data or management screens.\n\n"
        "Resort owners — Report recipients only (WhatsApp EOD when RP03 is live). Not app users in v2.\n\n"
        "System — Application automation: rollups, alerts, reconciliation, photo storage, scheduled jobs."
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Rozana Ops · Flow Inventory v2 · Generated from live codebase")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fr.italic = True

    return doc


def main() -> None:
    doc = build()
    OUT_DOCS.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCS)
    print(f"Saved: {OUT_DOCS}")
    if OUT_DOWNLOADS.parent.exists():
        doc.save(OUT_DOWNLOADS)
        print(f"Saved: {OUT_DOWNLOADS}")


if __name__ == "__main__":
    main()
