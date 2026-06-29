"""Generate Rozana Resorts go-live installation roadmap (Word)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Rozana-Go-Live-Roadmap.docx"
FOREST = RGBColor(0x1B, 0x43, 0x32)


def shade(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = FOREST


def para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade(c, "D8E8DF")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = val
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Rozana Resorts — Go-Live Roadmap")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = FOREST

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "Resort PC installation · Biometric enrolment · Ops web app · Daily WhatsApp report"
    )
    sr.font.size = Pt(11)
    sr.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Prepared for: Rozana Resorts (Ooty Skyview go-live)\nJune 2026")
    mr.font.size = Pt(10)

    doc.add_paragraph()

    heading(doc, "1. What you are installing", 1)
    para(
        doc,
        "Rozana is three connected pieces on one Supabase cloud database. "
        "This roadmap assumes development is complete and covers resort-side installation only.",
    )
    table(
        doc,
        ["Component", "Where it runs", "Purpose"],
        [
            ["Rozana Ops (web app)", "Browser on any PC/phone", "Manager & owner: attendance, HK, leave, reports"],
            ["Attendance kiosk (Windows)", "Dedicated PC at front desk", "Staff punch IN/OUT with fingerprint"],
            ["WhatsApp EOD sender (Baileys)", "Office PC or small VPS (always on)", "Daily report to owner at 6 PM IST"],
            ["Supabase (cloud)", "Already live", "Database, photos, Edge Functions — no install at resort"],
        ],
    )

    heading(doc, "2. Before you travel to the resort", 1)
    heading(doc, "2.1 Hardware checklist", 2)
    table(
        doc,
        ["Item", "Specification", "Notes"],
        [
            ["Resort kiosk PC", "Windows 10/11, 8 GB RAM, SSD", "Always on; disable sleep"],
            ["USB fingerprint reader", "Mantra / SecuGen / Morpho with Windows SDK", "Must support 1:N identify + enrolment"],
            ["UPS", "Recommended", "Prevents corrupt punches during power cut"],
            ["Stable internet", "Wi‑Fi or Ethernet", "Kiosk queues punches offline if dropped"],
            ["Dedicated WhatsApp SIM (optional)", "For EOD reports only", "Linked via QR to Baileys sender PC"],
        ],
    )

    heading(doc, "2.2 Accounts & credentials to collect", 2)
    bullets(
        doc,
        [
            "Supabase project URL and anon key (same as Rozana Ops .env.local)",
            "Supabase service_role key (Baileys sender only — keep secret)",
            "Owner WhatsApp number (international digits, e.g. 9198XXXXXXXX)",
            "Operations manager login (created in Rozana Ops)",
            "Kiosk ADMIN_PIN (change from default 1234 before go-live)",
        ],
    )

    heading(doc, "2.3 Software to copy to a USB stick", 2)
    bullets(
        doc,
        [
            "Flutter Windows release build: rozana_attendance_kiosk.exe + .env",
            "Fingerprint reader Windows drivers + SDK installer from vendor",
            "Node.js 18+ installer (for WhatsApp sender PC, if separate)",
            "Folder: Rozanaretreats/whatsapp-baileys (with .env configured)",
            "This document",
        ],
    )

    heading(doc, "3. Day 1 — Cloud database (one-time)", 1)
    para(doc, "Do this once in Supabase Dashboard (project oqksrcruypdmsggorrge) or confirm already done:")
    numbered(
        doc,
        [
            "SQL Editor → run rozana_attendance_kiosk/supabase/kiosk_migrations.sql",
            "Confirm tables: attendance_punches, fingerprint_templates, report_send_log",
            "Confirm Realtime on attendance_punches (see kiosk_migrations.sql bottom)",
            "Add all staff in Rozana Ops → Attendance → Add staff (with HK logins if needed)",
            "Add rooms / HK tasks for the property",
        ],
    )

    heading(doc, "4. Day 2 — Install Rozana Ops (web app)", 1)
    heading(doc, "4.1 Production hosting (recommended)", 2)
    bullets(
        doc,
        [
            "Deploy Rozanaretreats/app to Netlify, Vercel, or similar",
            "Set environment variables: VITE_SUPABASE_URL, VITE_SUPABASE_ANON_KEY",
            "Set VITE_ALLOW_MANUAL_PUNCHES=false when kiosks are live",
            "Managers open the URL in Chrome — bookmark on their phones",
        ],
    )

    heading(doc, "4.2 Local testing at resort (temporary)", 2)
    numbered(
        doc,
        [
            "Install Node.js on a laptop",
            "Copy Rozanaretreats/app folder",
            "Create .env.local with Supabase URL + anon key",
            "Run: npm install && npm run dev",
            "Open http://localhost:5173 — login as Firoz (manager) or Ruheed (owner)",
        ],
    )

    heading(doc, "5. Day 3 — Install fingerprint reader on resort PC", 1)
    numbered(
        doc,
        [
            "Install Windows updates; create dedicated Windows user (e.g. RozanaKiosk)",
            "Install vendor USB fingerprint driver + SDK",
            "Plug reader into USB; confirm device appears in Device Manager",
            "Disable sleep: Settings → System → Power → Never sleep when plugged in",
            "Set auto-login for RozanaKiosk user (optional, for 24/7 kiosk)",
        ],
    )
    para(
        doc,
        "Important: Until the Windows native plugin is wired to your exact reader SDK, "
        "you can run mock mode for training. For live fingerprints set FINGERPRINT_READER=native "
        "after SDK integration is complete.",
        bold=True,
    )

    heading(doc, "6. Day 4 — Install attendance kiosk app", 1)
    heading(doc, "6.1 Build release on dev machine", 2)
    numbered(
        doc,
        [
            "cd rozana_attendance_kiosk",
            "flutter build windows",
            "Output: build/windows/x64/runner/Release/",
            "Copy entire Release folder to resort PC, e.g. C:/Rozana/Kiosk/",
        ],
    )

    heading(doc, "6.2 Configure .env on resort PC", 2)
    para(doc, "Create C:/Rozana/Kiosk/.env with:")
    code = doc.add_paragraph()
    code.style = "No Spacing"
    for line in [
        "SUPABASE_URL=https://oqksrcruypdmsggorrge.supabase.co",
        "SUPABASE_ANON_KEY=<your anon key>",
        "PROPERTY_ID=ooty-skyview",
        "DEVICE_ID=kiosk-ooty-pc-01",
        "ADMIN_PIN=<choose 4–6 digits>",
        "FINGERPRINT_READER=native",
        "PUNCH_DEBOUNCE_SECONDS=30",
        "MIN_MATCH_SCORE=60",
    ]:
        r = code.add_run(line + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(9)

    heading(doc, "6.3 Auto-start at Windows logon", 2)
    numbered(
        doc,
        [
            "Open Task Scheduler → Create Task",
            "Trigger: At log on (RozanaKiosk user)",
            "Action: Start program → C:/Rozana/Kiosk/rozana_attendance_kiosk.exe",
            "Start in: C:/Rozana/Kiosk/",
            "Reboot PC and confirm kiosk opens full screen automatically",
        ],
    )

    heading(doc, "7. Biometric enrolment — step by step", 1)
    para(
        doc,
        "Fingerprints are enrolled on the kiosk only. Rozana Ops does not store templates. "
        "Each staff member is enrolled once before they can punch.",
    )

    heading(doc, "7.1 Who does what", 2)
    table(
        doc,
        ["Role", "Action"],
        [
            ["Operations manager (Firoz)", "Opens Admin on kiosk; enrols each staff fingerprint"],
            ["HK staff", "Only places finger on idle screen — no login"],
            ["Owner (Ruheed)", "Views attendance in Rozana Ops — does not enrol"],
        ],
    )

    heading(doc, "7.2 Enrol one staff member", 2)
    numbered(
        doc,
        [
            "Confirm staff exists in Rozana Ops (Attendance → staff list) for this property",
            "On kiosk idle screen, tap bottom label “Rozana Kiosk · device-id” five times quickly",
            "OR tap Admin button (visible in mock/desktop mode)",
            "Enter ADMIN_PIN from .env",
            "Select staff name from dropdown",
            "Tap “Capture fingerprint” — staff places finger on reader when prompted",
            "Wait for green “Enrolled successfully” message",
            "Repeat for each staff member (Arun, Vijay, Janaki, etc.)",
        ],
    )

    heading(doc, "7.3 Re-enrol or replace fingerprint", 2)
    bullets(
        doc,
        [
            "Same Admin flow — if already enrolled, app warns and allows capture again to replace",
            "Use after reader swap, poor quality scan, or staff finger injury",
        ],
    )

    heading(doc, "7.4 Test punch after enrolment", 2)
    numbered(
        doc,
        [
            "Return to kiosk home (back from Admin)",
            "Staff places finger on reader",
            "First punch of the day → CHECKED IN (green banner)",
            "Second punch → CHECKED OUT",
            "Open Rozana Ops → Attendance — punch appears within 1–2 seconds (Kiosk badge)",
        ],
    )

    heading(doc, "8. Daily operations (after go-live)", 1)
    table(
        doc,
        ["Time", "What happens"],
        [
            ["All day", "Staff punch IN/OUT at kiosk; manager uses Ops for HK & leave"],
            ["Real-time", "Ops Attendance updates live via Supabase"],
            ["6:00 PM IST", "Baileys sends EOD WhatsApp report to owner"],
            ["Anytime", "Owner logs into Ops → Reports / Housekeeping for proof photos"],
        ],
    )

    heading(doc, "9. WhatsApp daily report (Baileys)", 1)
    numbered(
        doc,
        [
            "On always-on PC: cd Rozanaretreats/whatsapp-baileys",
            "Edit .env: OWNER_WHATSAPP_NUMBERS, SUPABASE_ANON_KEY, EOD_CRON=0 18 * * *",
            "Run start.cmd → scan QR once with sender phone (Linked devices)",
            "Test: send-now.cmd",
            "Leave start.cmd running; cron sends daily automatically",
            "Owner views proof photos in Ops — WhatsApp message has no photo links",
        ],
    )

    heading(doc, "10. Go-live checklist", 1)
    bullets(
        doc,
        [
            "☐ kiosk_migrations.sql applied on Supabase",
            "☐ All staff added in Rozana Ops",
            "☐ Fingerprint reader drivers installed on resort PC",
            "☐ Kiosk .env configured (PROPERTY_ID, DEVICE_ID, ADMIN_PIN)",
            "☐ Every staff member fingerprint enrolled",
            "☐ Test IN + OUT punch; visible in Ops Attendance",
            "☐ VITE_ALLOW_MANUAL_PUNCHES=false in production Ops",
            "☐ Baileys connected; test EOD received on owner number",
            "☐ Task Scheduler auto-starts kiosk on boot",
            "☐ Manager trained: Admin PIN, enrolment, Ops HK workflow",
        ],
    )

    heading(doc, "11. Troubleshooting", 1)
    table(
        doc,
        ["Problem", "Fix"],
        [
            ["Kiosk says “Finger not recognized”", "Re-enrol in Admin; check reader cable"],
            ["Punches not in Ops", "Check internet; verify Supabase URL in .env"],
            ["Ops shows old data", "Hard refresh; check Supabase connection"],
            ["WhatsApp not sent", "Is start.cmd running? Run send-now.cmd to test"],
            ["QR expired (Baileys)", "Restart start.cmd; scan new QR"],
            ["PC rebooted — kiosk not open", "Check Task Scheduler path to .exe"],
        ],
    )

    heading(doc, "12. Support contacts & repo", 1)
    bullets(
        doc,
        [
            "GitHub: https://github.com/Rozanaretreats/Rozana-retreats",
            "Supabase project: oqksrcruypdmsggorrge",
            "Kiosk docs: rozana_attendance_kiosk/docs/PC_KIOSK_DEPLOYMENT.md",
            "When reader model is known: share brand + SDK link for Windows plugin completion",
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— End of Rozana Go-Live Roadmap —")
    fr.italic = True
    fr.font.size = Pt(10)

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
