"""Print-friendly Rozana handover checklist — Netlify + kiosk + mobile web (no codebase to client)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Rozana-Handover-Print-Checklist.docx"
FOREST = RGBColor(0x1B, 0x43, 0x32)
BOX = "\u2610"


def shade(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = FOREST


def todo(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{BOX}  {text}")
    r.font.size = Pt(12)


def para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(11)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade(c, "D8E8DF")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = val
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("ROZANA RETREATS — HANDOVER CHECKLIST")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = FOREST

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Ooty Skyview · Print & tick · You keep code · Client gets links & logins only")
    sr.font.size = Pt(11)

    doc.add_paragraph()

    # --- ARCHITECTURE ---
    heading(doc, "How it works (3 pieces — not one app)", 1)
    table(
        doc,
        ["Piece", "Who uses it", "Device", "What you give client"],
        [
            [
                "1. Rozana Ops (web)",
                "Ruheed, Firoz, HK staff",
                "Desktop + mobile browser",
                "One Netlify URL + logins",
            ],
            [
                "2. Attendance kiosk",
                "All staff (punch only)",
                "Windows PC at desk",
                "Nothing to install themselves — you set up PC",
            ],
            [
                "3. WhatsApp EOD",
                "Ruheed (owner)",
                "His phone",
                "Automatic message — no link",
            ],
        ],
    )

    para(
        doc,
        "You do NOT give the GitHub repo or source code. You deploy once; clients use links.",
        bold=True,
    )

    heading(doc, "Do HK staff need a separate mobile app?", 2)
    para(
        doc,
        "No. Rozana Ops already has a mobile staff portal (My Tasks, Attendance, Leave) "
        "with bottom navigation on phones. HK staff open the same Netlify link in Chrome/Safari "
        "and bookmark it (Add to Home Screen). No App Store build required for go-live.",
    )
    todo(doc, "Optional later: PWA icon / manifest polish — not blocking handover")

    heading(doc, "Netlify — yes, link is enough for Ops", 2)
    para(
        doc,
        "After you deploy Rozana Ops to Netlify with Supabase env vars baked in at build time, "
        "clients only need the URL. They never need npm, Node, or your laptop.",
    )

    doc.add_page_break()

    # --- BEFORE HANDOVER ---
    heading(doc, "SECTION 1 — YOU do this before handover (office)", 1)
    para(doc, "Tick each box. Client should not see any of this work.", bold=True)

    heading(doc, "1.1 Supabase (cloud database)", 2)
    todo(doc, "Project live: oqksrcruypdmsggorrge")
    todo(doc, "kiosk_migrations.sql applied")
    todo(doc, "All Ooty staff added (names, roles, property ooty-skyview)")
    todo(doc, "HK staff app logins created (email + password per person)")
    todo(doc, "Rooms / HK tasks set up for Ooty")
    todo(doc, "Security migration 011 applied (hashed passwords)")
    todo(doc, "Deploy latest send-owner-eod-report Edge Function")
    todo(doc, "Keep service_role key SECRET — never give to client")

    heading(doc, "1.2 Deploy Rozana Ops to Netlify", 2)
    todo(doc, "Netlify site connected to repo OR manual upload of build folder")
    todo(doc, "Base directory: Rozanaretreats/app")
    todo(doc, "Build command: npm run build")
    todo(doc, "Publish directory: dist")
    todo(doc, "Environment variables on Netlify:")
    todo(doc, "    VITE_SUPABASE_URL = https://oqksrcruypdmsggorrge.supabase.co")
    todo(doc, "    VITE_SUPABASE_ANON_KEY = (anon key from Supabase → API)")
    todo(doc, "    VITE_ALLOW_MANUAL_PUNCHES = false  (true only if kiosk not live yet)")
    todo(doc, "Add SPA redirect so refresh works: public/_redirects file with: /* /index.html 200")
    todo(doc, "Deploy succeeds — open production URL in browser")
    todo(doc, "Test login: Ruheed (owner), Firoz (manager), one HK staff on phone")

    heading(doc, "1.3 Build attendance kiosk (Windows — resort PC only)", 2)
    note_lines = [
        "This is NOT on Netlify. It is a Windows program on a PC at the property.",
        "USB fingerprint: native SDK still TODO — use mock/click for training if reader not wired.",
    ]
    for n in note_lines:
        p = doc.add_paragraph()
        r = p.add_run(n)
        r.italic = True
        r.font.size = Pt(10)

    todo(doc, "flutter build windows → copy Release folder to USB")
    todo(doc, "Prepare .env: PROPERTY_ID=ooty-skyview, DEVICE_ID=kiosk-ooty-pc-01, ADMIN_PIN=(secret)")
    todo(doc, "FINGERPRINT_READER=mock OR native (when reader SDK ready)")

    heading(doc, "1.4 WhatsApp daily report (Baileys — you run, not client)", 2)
    todo(doc, "whatsapp-baileys .env: OWNER_WHATSAPP_NUMBERS = Ruheed")
    todo(doc, "Test start.cmd + send-now.cmd once")
    todo(doc, "Plan: always-on PC at resort or your office runs start.cmd daily")

    heading(doc, "1.5 Write client handover sheet (fill in blanks)", 2)
    handover = doc.add_paragraph()
    for line in [
        "Rozana Ops URL: https://_________________________.netlify.app",
        "",
        "Ruheed (owner)     email: _______________  password: _______________",
        "Firoz (manager)    email: _______________  password: _______________",
        "HK staff logins    (list on separate page — one per person)",
        "",
        "Kiosk ADMIN_PIN: (Firoz only — not HK staff)",
        "Support contact: _______________",
    ]:
        r = handover.add_run(line + "\n")
        r.font.size = Pt(11)

    doc.add_page_break()

    # --- ON SITE ---
    heading(doc, "SECTION 2 — Ooty visit (one day — make it running)", 1)

    heading(doc, "2.1 Resort PC + kiosk (1–2 hours)", 2)
    todo(doc, "PC on power + internet; disable sleep")
    todo(doc, "Install fingerprint reader driver (if hardware ready)")
    todo(doc, "Copy kiosk to C:/Rozana/Kiosk/ + .env")
    todo(doc, "Run rozana_attendance_kiosk.exe — full screen")
    todo(doc, "Task Scheduler: auto-start kiosk at Windows logon")

    heading(doc, "2.2 Enrol every staff fingerprint on kiosk", 2)
    todo(doc, "Admin: tap footer 5× or Admin button → enter ADMIN_PIN")
    todo(doc, "For each staff: select name → Capture fingerprint → success")
    todo(doc, "Test: punch IN → see punch in Ops (phone) within 2 seconds")
    todo(doc, "Test: punch OUT")

    heading(doc, "2.3 Give everyone the Netlify link", 2)
    table(
        doc,
        ["Person", "Device", "URL", "Login"],
        [
            ["Ruheed", "Desktop/laptop", "Same Netlify URL", "Owner account"],
            ["Firoz", "Desktop + phone", "Same Netlify URL", "Manager account"],
            ["HK staff", "Mobile phone", "Same Netlify URL — bookmark", "Staff email/password"],
        ],
    )
    todo(doc, "HK phones: Chrome → open URL → Add to Home Screen (optional)")
    todo(doc, "Show HK: My Tasks → start room → checklist → verification photos")

    heading(doc, "2.4 WhatsApp EOD on site", 2)
    todo(doc, "Run start.cmd on always-on PC → scan QR once")
    todo(doc, "send-now.cmd → Ruheed receives report on phone")
    todo(doc, "Leave process running for 6 PM IST auto-send")

    heading(doc, "2.5 Training (same day)", 2)
    todo(doc, "Firoz: Ops on desktop — HK assign, attendance, leave, Reports")
    todo(doc, "Firoz: kiosk Admin PIN only — re-enrol if needed")
    todo(doc, "HK: finger punch at kiosk + My Tasks on phone")
    todo(doc, "Ruheed: Reports + proof photos in Ops; WhatsApp is summary only")

    heading(doc, "2.6 Sign-off before you leave Ooty", 2)
    todo(doc, "All staff enrolled + punch test OK")
    todo(doc, "One HK room completed with photos visible in Ops")
    todo(doc, "Netlify URL works on Firoz phone + Ruheed laptop")
    todo(doc, "Kiosk survives reboot")
    todo(doc, "WhatsApp test report received")
    todo(doc, "Handover sheet signed / photos of setup for your records")

    doc.add_page_break()

    # --- CLIENT CHEAT SHEET ---
    heading(doc, "SECTION 3 — Give client (print this page)", 1)
    para(doc, "Rozana Ops — open in browser (bookmark this):", bold=True)
    p = doc.add_paragraph()
    p.add_run("https://________________________________.netlify.app").font.size = Pt(14)

    para(doc, "Daily use:", bold=True)
    bullets_client = [
        "HK staff: punch finger at front-desk PC when arriving and leaving",
        "HK staff: phone → Rozana Ops → Tasks → clean room → take verification photos",
        "Firoz: phone/laptop → assign rooms, approve leave, check attendance",
        "Ruheed: laptop → Reports & Housekeeping for proof photos",
        "Ruheed: WhatsApp summary every evening ~6 PM (automatic)",
    ]
    for b in bullets_client:
        doc.add_paragraph(b, style="List Bullet")

    para(doc, "Problems — call: _______________________", bold=True)

    heading(doc, "SECTION 4 — What you NEVER give client", 1)
    for item in [
        "GitHub / source code",
        "Supabase service_role key",
        "Database password",
        "whatsapp-baileys auth_info folder",
        "Your personal Supabase dashboard access (optional: add Firoz view-only later)",
    ]:
        todo(doc, f"Do not share: {item}")

    heading(doc, "SECTION 5 — Quick troubleshooting", 1)
    table(
        doc,
        ["Issue", "Fix"],
        [
            ["Ops page blank after refresh", "Netlify SPA redirect missing"],
            ["HK cannot login", "Check staff email/password in Supabase staff_logins"],
            ["Punch not in Ops", "Kiosk internet; check .env on PC"],
            ["Finger not recognized", "Admin re-enrol on kiosk"],
            ["No evening WhatsApp", "Baileys PC must stay on; restart start.cmd"],
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Rozana Handover Print Checklist — Ooty Skyview")
    fr.italic = True
    fr.font.size = Pt(10)

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
