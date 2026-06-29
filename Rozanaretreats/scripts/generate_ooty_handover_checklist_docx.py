"""Generate Ooty 1-day handover & implementation checklist (Word, todo format)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Rozana-Ooty-1-Day-Handover-Checklist.docx"
FOREST = RGBColor(0x1B, 0x43, 0x32)
BOX = "\u2610"  # ☐ empty checkbox


def shade(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = FOREST


def todo(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{BOX} {text}")
    run.font.size = Pt(11)


def note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade(c, "D8E8DF")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = val
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def build() -> Document:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Rozana — Ooty Go-Live in 1 Day")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = FOREST

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Pre-handover checklist · On-site implementation · Biometrics · Staff training")
    sr.font.size = Pt(11)
    sr.italic = True

    doc.add_paragraph()
    note(
        doc,
        "Based on project scan (June 2026). Property: Ooty Skyview (ooty-skyview). "
        "Supabase: oqksrcruypdmsggorrge. GitHub: Rozanaretreats/Rozana-retreats.",
    )

    heading(doc, "What is ready vs what you must do on site", 1)
    table(
        doc,
        ["Area", "Status", "Your action"],
        [
            ["Rozana Ops web app", "Ready", "Deploy URL + logins for Firoz & Ruheed"],
            ["Kiosk app (Windows)", "Ready (mock mode)", "Native USB reader SDK still TODO — see Section 2"],
            ["Supabase data layer", "Live", "Confirm migrations + staff/rooms loaded"],
            ["Live attendance sync", "Ready", "Realtime on attendance_punches"],
            ["EOD WhatsApp (Baileys)", "Ready", "start.cmd + QR on always-on PC"],
            ["EOD Edge Function", "Deployed", "Deploy latest version (attendance time fix)"],
            ["Windows fingerprint plugin", "NOT done", "Only if USB reader purchased + SDK integrated"],
            ["Ops Reports WhatsApp button", "Disabled UI", "Use Baileys sender — not this button"],
        ],
    )

    # --- BEFORE HANDOVER ---
    heading(doc, "PART A — Before you travel (office / home)", 1)
    note(doc, "Complete these before leaving for Ooty. Do not hand over until all critical items are checked.")

    heading(doc, "A1. Cloud & credentials", 2)
    todo(doc, "Supabase project accessible — save URL + anon key + service_role key (secret)")
    todo(doc, "Run kiosk_migrations.sql in SQL Editor (if not already applied)")
    todo(doc, "Confirm tables: staff, rooms, housekeeping_tasks, attendance_punches, fingerprint_templates, report_send_log")
    todo(doc, "Deploy latest send-owner-eod-report Edge Function (attendance IN–OUT fix)")
    todo(doc, "Run security migration 011 if staff passwords are still plain text")

    heading(doc, "A2. Rozana Ops (manager + owner app)", 2)
    todo(doc, "Deploy Rozanaretreats/app to production host (Netlify/Vercel) — not localhost")
    todo(doc, "Set VITE_SUPABASE_URL and VITE_SUPABASE_ANON_KEY on host")
    todo(doc, "Set VITE_ALLOW_MANUAL_PUNCHES=false for go-live (or true only if no reader yet)")
    todo(doc, "Create Firoz login (operations manager, Ooty only)")
    todo(doc, "Confirm Ruheed owner login works (ruheed@rozana.com)")
    todo(doc, "Add all Ooty staff in Attendance → Add staff (with HK app logins)")
    todo(doc, "Set up room list / HK tasks for Ooty property")
    todo(doc, "Test: login, attendance page, HK task flow, photo upload, Reports page")

    heading(doc, "A3. Kiosk build on USB stick", 2)
    todo(doc, "On dev PC: cd rozana_attendance_kiosk → flutter build windows")
    todo(doc, "Copy Release folder to USB: rozana_attendance_kiosk.exe + flutter DLLs + data folder")
    todo(doc, "Create .env for Ooty PC (see A4) on USB")
    todo(doc, "Copy whatsapp-baileys folder + configured .env to USB")
    todo(doc, "Copy fingerprint reader Windows driver installer (if hardware bought)")
    todo(doc, "Copy this checklist + Rozana-Go-Live-Roadmap.docx")

    heading(doc, "A4. Kiosk .env template (fill before trip)", 2)
    code = doc.add_paragraph()
    for line in [
        "SUPABASE_URL=https://oqksrcruypdmsggorrge.supabase.co",
        "SUPABASE_ANON_KEY=<anon key>",
        "PROPERTY_ID=ooty-skyview",
        "DEVICE_ID=kiosk-ooty-pc-01",
        "ADMIN_PIN=<choose — tell only Firoz>",
        "FINGERPRINT_READER=mock   ← change to native when reader SDK wired",
        "PUNCH_DEBOUNCE_SECONDS=30",
        "MIN_MATCH_SCORE=60",
    ]:
        r = code.add_run(line + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(9)

    heading(doc, "A5. WhatsApp EOD sender", 2)
    todo(doc, "whatsapp-baileys/.env: OWNER_WHATSAPP_NUMBERS = Ruheed (or test number)")
    todo(doc, "npm install completed in whatsapp-baileys folder")
    todo(doc, "Test send-now.cmd once from office (optional dry run)")

    heading(doc, "A6. Hardware decision (critical)", 2)
    note(
        doc,
        "Project scan: Windows USB fingerprint SDK is NOT integrated yet (TODO in native_fingerprint_reader.dart). "
        "For real fingerprints at Ooty you need EITHER: (1) reader + SDK integrated before trip, OR "
        "(2) go-live Day 1 in mock/click mode for training, reader integrated later.",
    )
    todo(doc, "Confirm: USB reader purchased? Brand/model? Windows SDK installer obtained?")
    todo(doc, "If NO reader yet: plan mock mode training; schedule SDK integration date")
    todo(doc, "If YES reader: integrate Windows plugin BEFORE Ooty or accept mock-only on Day 1")

    heading(doc, "A7. Handover gate — do not leave until", 2)
    todo(doc, "Production Ops URL opens on phone and laptop")
    todo(doc, "All Ooty staff exist in database")
    todo(doc, "Kiosk Release build runs on your laptop (mock test)")
    todo(doc, "Baileys test message received on owner number")
    todo(doc, "ADMIN_PIN chosen and documented for Firoz (not default 1234)")

    # --- ONE DAY OOTY ---
    heading(doc, "PART B — Ooty: single-day implementation (your visit)", 1)
    note(doc, "Suggested order. Adjust times to your travel schedule.")

    heading(doc, "B1. Morning — PC & network (1–2 hours)", 2)
    todo(doc, "Unpack resort PC; connect power + UPS; connect internet")
    todo(doc, "Disable sleep/hibernate; set never turn off display while plugged in")
    todo(doc, "Install Windows updates if needed; create user RozanaKiosk (optional)")
    todo(doc, "Install USB fingerprint reader driver (if hardware present)")
    todo(doc, "Copy kiosk Release folder to C:/Rozana/Kiosk/")
    todo(doc, "Place .env in C:/Rozana/Kiosk/")
    todo(doc, "Double-click rozana_attendance_kiosk.exe — confirm app opens full screen")
    todo(doc, "Confirm property name shows Ooty Skyview on kiosk screen")

    heading(doc, "B2. Mid-morning — Rozana Ops on manager phone/laptop (30 min)", 2)
    todo(doc, "Open production Ops URL; bookmark on Firoz phone")
    todo(doc, "Firoz logs in; confirm Ooty property only (not Kannur)")
    todo(doc, "Walk through: Attendance, Leave, Housekeeping, Reports (Ruheed if present)")
    todo(doc, "Assign today’s HK rooms to staff in Housekeeping page")

    heading(doc, "B3. Biometric enrolment — every staff member (1–2 hours)", 2)
    note(doc, "Enrolment is on the KIOSK only. Ops app does not capture fingerprints.")

    heading(doc, "How to open Admin on kiosk", 3)
    todo(doc, "Tap “Rozana Kiosk · device-id” at bottom 5 times quickly, OR use Admin button (desktop mock mode)")
    todo(doc, "Enter ADMIN_PIN")
    todo(doc, "Select staff name from dropdown")
    todo(doc, "Tap Capture fingerprint — staff places finger when prompted (or click in mock mode)")
    todo(doc, "Wait for green success message")
    todo(doc, "Repeat for EVERY active staff member (Arun, Vijay, Janaki, …)")
    todo(doc, "Write enrolment log on paper: name ✓ / failed / re-try")

    heading(doc, "After each enrolment — test punch", 3)
    todo(doc, "Return to kiosk home screen")
    todo(doc, "Staff punches IN — green CHECKED IN banner")
    todo(doc, "Open Ops → Attendance on phone — punch appears within 2 sec (Kiosk badge)")
    todo(doc, "Staff punches OUT — confirm second punch on Ops")
    todo(doc, "If not recognized: re-enrol in Admin")

    heading(doc, "B4. Afternoon — WhatsApp daily report (30 min)", 2)
    todo(doc, "On always-on PC: install Node.js if missing")
    todo(doc, "Copy whatsapp-baileys folder; verify .env")
    todo(doc, "Run start.cmd → scan QR with sender phone (Linked devices)")
    todo(doc, "Wait for [baileys] Connected")
    todo(doc, "Run send-now.cmd in second window — Ruheed receives EOD report")
    todo(doc, "Confirm report format: attendance, HK, outstanding, no photo links")
    todo(doc, "Leave start.cmd running for 6 PM IST auto-send")

    heading(doc, "B5. Auto-start kiosk on boot (20 min)", 2)
    todo(doc, "Task Scheduler → At log on → run C:/Rozana/Kiosk/rozana_attendance_kiosk.exe")
    todo(doc, "Start in: C:/Rozana/Kiosk/")
    todo(doc, "Reboot PC once — confirm kiosk opens automatically")

    heading(doc, "B6. Staff training (1 hour)", 2)

    heading(doc, "Train Firoz (operations manager)", 3)
    todo(doc, "Ops login on his phone — bookmark URL")
    todo(doc, "Housekeeping: assign rooms, verify photos, mark done")
    todo(doc, "Attendance: read punches; mark absence if needed")
    todo(doc, "Leave: approve/reject requests")
    todo(doc, "Kiosk Admin PIN — how to re-enrol a finger if staff changes")
    todo(doc, "Do NOT share ADMIN_PIN with HK staff")

    heading(doc, "Train HK staff (Arun, Vijay, Janaki, …)", 3)
    todo(doc, "Punch: place finger on reader at start and end of shift — no buttons")
    todo(doc, "HK tasks: open Ops on phone → My Tasks → start cleaning → checklist → verification photos")
    todo(doc, "Practice one full room clean with photos while you watch")
    todo(doc, "Explain: owner gets WhatsApp summary at 6 PM — they cannot edit it")

    heading(doc, "Train Ruheed (owner) — 15 min if present", 3)
    todo(doc, "Login as owner on laptop")
    todo(doc, "Reports + Housekeeping for proof photos")
    todo(doc, "Daily WhatsApp report at 6 PM — proof is in Ops, not in WhatsApp")

    heading(doc, "B7. End of day — sign-off before you leave Ooty", 2)
    todo(doc, "Every staff: enrolled + successful IN punch tested")
    todo(doc, "At least one HK room completed with verification photos in Ops")
    todo(doc, "Firoz can assign HK and approve leave without your help")
    todo(doc, "Kiosk survives PC reboot (Task Scheduler)")
    todo(doc, "WhatsApp EOD received or send-now tested")
    todo(doc, "Write handover note: PC location, ADMIN_PIN holder, sender phone for WhatsApp, your contact")

    # --- AFTER ---
    heading(doc, "PART C — After handover (remote support)", 1)
    todo(doc, "Day 2: call Firoz — confirm morning punches appeared in Ops")
    todo(doc, "Day 2: confirm 6 PM WhatsApp arrived on Ruheed phone")
    todo(doc, "Week 1: monitor report_send_log in Supabase for failed sends")
    todo(doc, "When USB reader SDK ready: set FINGERPRINT_READER=native, re-enrol all staff")
    todo(doc, "Rotate ADMIN_PIN if shared during training")
    todo(doc, "Keep Baileys PC online daily (or move to office VPS)")

    heading(doc, "PART D — Quick troubleshooting (leave with Firoz)", 1)
    table(
        doc,
        ["Problem", "Fix"],
        [
            ["Finger not recognized", "Admin → re-enrol; clean reader surface"],
            ["Punch not in Ops", "Check internet on PC; verify .env Supabase URL"],
            ["Kiosk not open after reboot", "Check Task Scheduler path"],
            ["No WhatsApp at 6 PM", "Is start.cmd running? Wi‑Fi on sender PC?"],
            ["Baileys disconnected", "Restart start.cmd; re-scan QR if logged out"],
            ["Wrong IN–OUT times in report", "Redeploy latest Edge Function"],
        ],
    )

    heading(doc, "PART E — File locations in project", 1)
    bullets = [
        "Ops app: Rozanaretreats/app/",
        "Kiosk: rozana_attendance_kiosk/",
        "Kiosk SQL: rozana_attendance_kiosk/supabase/kiosk_migrations.sql",
        "WhatsApp sender: Rozanaretreats/whatsapp-baileys/ (start.cmd, send-now.cmd)",
        "EOD function: Rozanaretreats/supabase/functions/send-owner-eod-report/",
        "PC deployment notes: rozana_attendance_kiosk/docs/PC_KIOSK_DEPLOYMENT.md",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(10)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— End of Ooty 1-Day Handover Checklist —")
    fr.italic = True
    fr.font.size = Pt(10)

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
